import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except Exception as e:
    print("python-docx is required. Install with: pip install python-docx", file=sys.stderr)
    raise


def is_bullet(line: str) -> bool:
    ls = line.lstrip()
    return ls.startswith("- ") or ls.startswith("* ")


def is_heading(line: str) -> int:
    i = 0
    while i < len(line) and line[i] == '#':
        i += 1
    if i > 0 and (i < len(line) and line[i] == ' '):
        return min(i, 3)  # map up to Heading 3
    return 0


def add_code_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)


def md_to_docx(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding='utf-8')
    lines = text.splitlines()
    doc = Document()
    # Title from first H1 if present, else filename
    title = None
    for ln in lines:
        level = is_heading(ln)
        if level == 1:
            title = ln.lstrip('#').strip()
            break
    if title is None:
        title = md_path.stem
    doc.add_heading(title, level=0)

    in_code = False
    for raw in lines:
        line = raw.rstrip('\n')
        if line.strip().startswith('```'):
            in_code = not in_code
            if in_code:
                # start code block marker ignored
                pass
            else:
                # end code block marker ignored
                pass
            continue

        if in_code:
            add_code_paragraph(doc, line)
            continue

        # Headings
        level = is_heading(line)
        if level:
            text = line[level + 1:].strip() if line[level:level+1] == ' ' else line.lstrip('#').strip()
            doc.add_heading(text, level=level)
            continue

        # Bullets
        if is_bullet(line):
            content = line.lstrip()[2:].strip()
            p = doc.add_paragraph(content)
            try:
                p.style = 'List Bullet'
            except Exception:
                # Fallback: normal paragraph
                pass
            continue

        # Blank line
        if not line.strip():
            doc.add_paragraph("")
            continue

        # Regular paragraph
        doc.add_paragraph(line)

    doc.save(str(docx_path))


def main():
    if len(sys.argv) < 3:
        print("Usage: python md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    md = Path(sys.argv[1])
    out = Path(sys.argv[2])
    md_to_docx(md, out)
    print(f"Wrote {out}")


if __name__ == '__main__':
    main()

